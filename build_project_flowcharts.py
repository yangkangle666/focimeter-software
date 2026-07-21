from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\yangkangle\OneDrive\Desktop\焦度计")
OUT_DIR = ROOT / "outputs"
IMG_DIR = OUT_DIR / "流程图图片"
DOCX_PATH = OUT_DIR / "焦度计项目三类流程图.docx"
PDF_PATH = OUT_DIR / "焦度计项目三类流程图.pdf"


FONT_PATHS = [
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\simsun.ttc",
]


def font(size, bold=False):
    path = next((p for p in FONT_PATHS if Path(p).exists()), None)
    if not path:
        return ImageFont.load_default()
    return ImageFont.truetype(path, size=size, index=0)


F_TITLE = font(48, True)
F_SUB = font(26)
F_NODE = font(25)
F_SMALL = font(21)
F_TAG = font(22)


COLORS = {
    "bg": "#F7FAFC",
    "ink": "#0B2545",
    "muted": "#5B677A",
    "line": "#34506B",
    "blue": "#DCEEFF",
    "blue_border": "#2F6FA5",
    "green": "#DEF7E5",
    "green_border": "#2F8F4E",
    "orange": "#FFEBCF",
    "orange_border": "#CC7A1A",
    "purple": "#EDE4FF",
    "purple_border": "#6D4BC3",
    "red": "#FFE2E2",
    "red_border": "#B84242",
    "gray": "#EEF2F6",
    "gray_border": "#7B8794",
}


def hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i + 2], 16) for i in (0, 2, 4))


def wrap_text(draw, text, fnt, max_width):
    lines = []
    for paragraph in text.split("\n"):
        current = ""
        for ch in paragraph:
            test = current + ch
            if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def center_text(draw, box, text, fnt, color=COLORS["ink"], line_gap=8):
    x, y, w, h = box
    lines = wrap_text(draw, text, fnt, w - 44)
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    cy = y + (h - total_h) / 2
    for line, lh in zip(lines, heights):
        tw = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x + (w - tw) / 2, cy), line, font=fnt, fill=hex_to_rgb(color))
        cy += lh + line_gap


def node(draw, box, text, fill, border, fnt=F_NODE):
    x, y, w, h = box
    draw.rounded_rectangle(
        (x, y, x + w, y + h),
        radius=24,
        fill=hex_to_rgb(fill),
        outline=hex_to_rgb(border),
        width=3,
    )
    center_text(draw, box, text, fnt)


def diamond(draw, box, text, fill, border):
    x, y, w, h = box
    points = [(x + w / 2, y), (x + w, y + h / 2), (x + w / 2, y + h), (x, y + h / 2)]
    draw.polygon(points, fill=hex_to_rgb(fill), outline=hex_to_rgb(border))
    for offset in range(1, 3):
        draw.line(points + [points[0]], fill=hex_to_rgb(border), width=offset)
    center_text(draw, box, text, F_NODE)


def arrow(draw, start, end, label=None, curved=False):
    sx, sy = start
    ex, ey = end
    color = hex_to_rgb(COLORS["line"])
    if curved:
        mid1 = (sx, (sy + ey) // 2)
        mid2 = (ex, (sy + ey) // 2)
        draw.line([start, mid1, mid2, end], fill=color, width=4, joint="curve")
    else:
        draw.line([start, end], fill=color, width=4)
    import math
    ang = math.atan2(ey - sy, ex - sx)
    head = 18
    p1 = (ex - head * math.cos(ang - math.pi / 6), ey - head * math.sin(ang - math.pi / 6))
    p2 = (ex - head * math.cos(ang + math.pi / 6), ey - head * math.sin(ang + math.pi / 6))
    draw.polygon([end, p1, p2], fill=color)
    if label:
        mx, my = (sx + ex) / 2, (sy + ey) / 2
        bbox = draw.textbbox((0, 0), label, font=F_SMALL)
        tw, th = bbox[2], bbox[3]
        draw.rounded_rectangle((mx - tw / 2 - 12, my - th / 2 - 9, mx + tw / 2 + 12, my + th / 2 + 9),
                               radius=10, fill=hex_to_rgb("#FFFFFF"), outline=hex_to_rgb("#CAD3DF"))
        draw.text((mx - tw / 2, my - th / 2 - 2), label, font=F_SMALL, fill=hex_to_rgb(COLORS["muted"]))


def poly_arrow(draw, points, label=None, label_pos=None):
    color = hex_to_rgb(COLORS["line"])
    draw.line(points, fill=color, width=4, joint="curve")
    import math
    (sx, sy), (ex, ey) = points[-2], points[-1]
    ang = math.atan2(ey - sy, ex - sx)
    head = 18
    p1 = (ex - head * math.cos(ang - math.pi / 6), ey - head * math.sin(ang - math.pi / 6))
    p2 = (ex - head * math.cos(ang + math.pi / 6), ey - head * math.sin(ang + math.pi / 6))
    draw.polygon([points[-1], p1, p2], fill=color)
    if label:
        if label_pos is None:
            label_pos = points[len(points) // 2]
        mx, my = label_pos
        bbox = draw.textbbox((0, 0), label, font=F_SMALL)
        tw, th = bbox[2], bbox[3]
        draw.rounded_rectangle((mx - tw / 2 - 12, my - th / 2 - 9, mx + tw / 2 + 12, my + th / 2 + 9),
                               radius=10, fill=hex_to_rgb("#FFFFFF"), outline=hex_to_rgb("#CAD3DF"))
        draw.text((mx - tw / 2, my - th / 2 - 2), label, font=F_SMALL, fill=hex_to_rgb(COLORS["muted"]))


def header(draw, title, subtitle):
    draw.text((90, 58), title, font=F_TITLE, fill=hex_to_rgb(COLORS["ink"]))
    draw.text((92, 122), subtitle, font=F_SUB, fill=hex_to_rgb(COLORS["muted"]))


def save_canvas(name):
    path = IMG_DIR / name
    return path


def project_flow():
    img = Image.new("RGB", (2200, 1400), hex_to_rgb(COLORS["bg"]))
    d = ImageDraw.Draw(img)
    header(d, "流程图一：整个项目的制作流程", "从答辩通过到本地系统、软件封装、设备联调和最终交付")

    boxes = {
        "a": (90, 230, 300, 105),
        "b": (470, 230, 300, 105),
        "c": (850, 230, 330, 105),
        "d": (1270, 230, 330, 105),
        "e": (1680, 230, 330, 105),
        "f": (1680, 500, 330, 105),
        "g": (1270, 500, 330, 105),
        "h": (850, 500, 330, 105),
        "i": (470, 500, 300, 105),
        "j": (90, 500, 300, 105),
        "k": (850, 770, 330, 105),
        "l": (1270, 770, 330, 105),
        "m": (1680, 770, 330, 105),
    }

    labels = {
        "a": "项目答辩通过\n明确接手资格",
        "b": "资料整理\n理解焦度计原理",
        "c": "确认边界\n软件为主 硬件配合",
        "d": "明确参数\n相机 光路 标定图",
        "e": "搭建本地系统\n先跑通完整流程",
        "f": "图像算法开发\n识别光斑和质心",
        "g": "标定与计算模型\n坐标 位移 度数",
        "h": "样本验证\n标准镜片对比",
        "i": "验证是否达标",
        "j": "算法优化\n参数修正",
        "k": "封装成软件\n界面 报告 配置",
        "l": "设备联调\n相机 硬件 软件",
        "m": "演示与交付\n文档 测试 成果",
    }
    fills = {
        "a": ("green", "green_border"), "b": ("blue", "blue_border"), "c": ("blue", "blue_border"),
        "d": ("orange", "orange_border"), "e": ("purple", "purple_border"), "f": ("purple", "purple_border"),
        "g": ("purple", "purple_border"), "h": ("orange", "orange_border"), "i": ("red", "red_border"),
        "j": ("red", "red_border"), "k": ("blue", "blue_border"), "l": ("orange", "orange_border"),
        "m": ("green", "green_border"),
    }
    for key, box in boxes.items():
        fill, border = fills[key]
        if key == "i":
            diamond(d, box, labels[key], COLORS[fill], COLORS[border])
        else:
            node(d, box, labels[key], COLORS[fill], COLORS[border])

    seq = ["a", "b", "c", "d", "e"]
    for s, e in zip(seq, seq[1:]):
        arrow(d, (boxes[s][0] + boxes[s][2], boxes[s][1] + boxes[s][3] // 2),
              (boxes[e][0], boxes[e][1] + boxes[e][3] // 2))
    arrow(d, (1845, 335), (1845, 500))
    for s, e in [("f", "g"), ("g", "h"), ("h", "i"), ("i", "j")]:
        arrow(d, (boxes[s][0], boxes[s][1] + boxes[s][3] // 2),
              (boxes[e][0] + boxes[e][2], boxes[e][1] + boxes[e][3] // 2), label="不达标" if s == "i" else None)
    poly_arrow(d, [(620, 605), (620, 705), (1015, 705), (1015, 770)], label="达标", label_pos=(820, 705))
    poly_arrow(d, [(240, 605), (240, 705), (1015, 705), (1015, 605)], label="重新验证", label_pos=(620, 705))
    arrow(d, (boxes["k"][0] + boxes["k"][2], boxes["k"][1] + boxes["k"][3] // 2),
          (boxes["l"][0], boxes["l"][1] + boxes["l"][3] // 2))
    arrow(d, (boxes["l"][0] + boxes["l"][2], boxes["l"][1] + boxes["l"][3] // 2),
          (boxes["m"][0], boxes["m"][1] + boxes["m"][3] // 2))

    d.text((96, 1190), "核心思路：先把项目理解清楚，再做本地系统验证算法；验证通过后再封装软件，最后和设备联调。", font=F_TAG, fill=hex_to_rgb(COLORS["muted"]))
    path = save_canvas("01_整个项目制作流程.png")
    img.save(path)
    return path


def software_flow():
    img = Image.new("RGB", (2200, 1400), hex_to_rgb(COLORS["bg"]))
    d = ImageDraw.Draw(img)
    header(d, "流程图二：软件部分的运行流程", "本地系统一次运行时，从导入图片到输出 S/C/A 的完整过程")

    xs = [100, 430, 760, 1090, 1420, 1750]
    y1, y2, y3 = 240, 510, 780
    w, h = 270, 100
    boxes = {
        "start": (xs[0], y1, w, h),
        "cfg": (xs[1], y1, w, h),
        "calib": (xs[2], y1, w, h),
        "pre1": (xs[3], y1, w, h),
        "spot1": (xs[4], y1, w, h),
        "coord": (xs[5], y1, w, h),
        "measure": (xs[5], y2, w, h),
        "pre2": (xs[4], y2, w, h),
        "spot2": (xs[3], y2, w, h),
        "check": (xs[2], y2, w, h),
        "shift": (xs[1], y2, w, h),
        "calc": (xs[0], y2, w, h),
        "out": (xs[0], y3, w, h),
        "log": (xs[1], y3, w, h),
        "err": (xs[3], y3, w, h),
        "redo": (xs[4], y3, w, h),
    }
    texts = {
        "start": "启动本地系统",
        "cfg": "读取配置\n像元 光路 阈值",
        "calib": "导入标定图",
        "pre1": "标定图预处理\n灰度 滤波 增强",
        "spot1": "识别光斑\n计算质心",
        "coord": "建立标定坐标系",
        "measure": "导入测量图",
        "pre2": "测量图预处理",
        "spot2": "识别测量光斑\n计算质心",
        "check": "识别是否成功",
        "shift": "坐标转换\n计算光斑位移",
        "calc": "判断镜片类型\n计算 S/C/A",
        "out": "输出结果\nS C A 质量状态",
        "log": "保存日志\n中间图和参数",
        "err": "异常提示\n说明失败原因",
        "redo": "重新采图\n或调整参数",
    }
    for key, box in boxes.items():
        if key in ["start", "out"]:
            fill, border = COLORS["green"], COLORS["green_border"]
        elif key in ["check", "err", "redo"]:
            fill, border = COLORS["red"], COLORS["red_border"]
        elif key in ["cfg", "log"]:
            fill, border = COLORS["orange"], COLORS["orange_border"]
        else:
            fill, border = COLORS["blue"], COLORS["blue_border"]
        if key == "check":
            diamond(d, box, texts[key], fill, border)
        else:
            node(d, box, texts[key], fill, border)

    for s, e in [("start", "cfg"), ("cfg", "calib"), ("calib", "pre1"), ("pre1", "spot1"), ("spot1", "coord")]:
        arrow(d, (boxes[s][0] + w, boxes[s][1] + h // 2), (boxes[e][0], boxes[e][1] + h // 2))
    arrow(d, (boxes["coord"][0] + w // 2, boxes["coord"][1] + h), (boxes["measure"][0] + w // 2, boxes["measure"][1]))
    for s, e in [("measure", "pre2"), ("pre2", "spot2"), ("spot2", "check"), ("check", "shift"), ("shift", "calc")]:
        arrow(d, (boxes[s][0], boxes[s][1] + h // 2), (boxes[e][0] + w, boxes[e][1] + h // 2), label="成功" if s == "check" else None)
    arrow(d, (boxes["calc"][0] + w // 2, boxes["calc"][1] + h), (boxes["out"][0] + w // 2, boxes["out"][1]))
    arrow(d, (boxes["out"][0] + w, boxes["out"][1] + h // 2), (boxes["log"][0], boxes["log"][1] + h // 2))
    arrow(d, (boxes["check"][0] + w // 2, boxes["check"][1] + h), (boxes["err"][0] + w // 2, boxes["err"][1]), label="失败")
    arrow(d, (boxes["err"][0] + w, boxes["err"][1] + h // 2), (boxes["redo"][0], boxes["redo"][1] + h // 2))
    arrow(d, (boxes["redo"][0] + w // 2, boxes["redo"][1]), (boxes["measure"][0] + w // 2, boxes["measure"][1] + h), label="再运行", curved=True)

    d.text((96, 1190), "核心输出：镜片参数 S/C/A、识别质量、中间图、运行日志。第一阶段先在本地系统中跑通。", font=F_TAG, fill=hex_to_rgb(COLORS["muted"]))
    path = save_canvas("02_软件部分运行流程.png")
    img.save(path)
    return path


def device_flow():
    img = Image.new("RGB", (2200, 1400), hex_to_rgb(COLORS["bg"]))
    d = ImageDraw.Draw(img)
    header(d, "流程图三：整个设备的运行流程", "一次镜片测量中，硬件成像和软件计算如何配合")

    boxes = {
        "power": (100, 240, 280, 100),
        "self": (450, 240, 280, 100),
        "ref": (800, 240, 300, 100),
        "place": (1170, 240, 300, 100),
        "capture": (1540, 240, 300, 100),
        "light": (1540, 500, 300, 100),
        "lens": (1170, 500, 300, 100),
        "mask": (800, 500, 300, 100),
        "camera": (450, 500, 280, 100),
        "software": (100, 500, 280, 100),
        "result": (100, 760, 280, 100),
        "quality": (450, 760, 280, 100),
        "report": (800, 760, 300, 100),
        "save": (1170, 760, 300, 100),
        "end": (1540, 760, 300, 100),
    }
    texts = {
        "power": "设备上电\n打开本地系统",
        "self": "硬件自检\n光源 相机 通信",
        "ref": "采集标定图\n建立参考基准",
        "place": "放入待测镜片\n固定位置",
        "capture": "开始测量\n采集测量图",
        "light": "光源发光\n准直光路形成光束",
        "lens": "光线经过镜片\n方向发生偏折",
        "mask": "哈特曼光阑\n形成多个采样光斑",
        "camera": "相机成像\n得到光斑图片",
        "software": "软件处理图片\n识别位移并计算",
        "result": "显示 S/C/A\n给出测量结果",
        "quality": "质量判断\n是否可信",
        "report": "生成记录\n结果和过程图",
        "save": "保存数据\n便于复查",
        "end": "结束测量\n更换下一片镜片",
    }
    for key, box in boxes.items():
        if key in ["power", "end"]:
            fill, border = COLORS["green"], COLORS["green_border"]
        elif key in ["quality"]:
            fill, border = COLORS["red"], COLORS["red_border"]
        elif key in ["software", "result", "report", "save"]:
            fill, border = COLORS["blue"], COLORS["blue_border"]
        elif key in ["light", "lens", "mask", "camera"]:
            fill, border = COLORS["purple"], COLORS["purple_border"]
        else:
            fill, border = COLORS["orange"], COLORS["orange_border"]
        if key == "quality":
            diamond(d, box, texts[key], fill, border)
        else:
            node(d, box, texts[key], fill, border)

    for s, e in [("power", "self"), ("self", "ref"), ("ref", "place"), ("place", "capture")]:
        arrow(d, (boxes[s][0] + boxes[s][2], boxes[s][1] + 50), (boxes[e][0], boxes[e][1] + 50))
    arrow(d, (1690, 340), (1690, 500))
    for s, e in [("light", "lens"), ("lens", "mask"), ("mask", "camera"), ("camera", "software")]:
        arrow(d, (boxes[s][0], boxes[s][1] + 50), (boxes[e][0] + boxes[e][2], boxes[e][1] + 50))
    arrow(d, (240, 600), (240, 760))
    for s, e in [("result", "quality"), ("quality", "report"), ("report", "save"), ("save", "end")]:
        arrow(d, (boxes[s][0] + boxes[s][2], boxes[s][1] + 50), (boxes[e][0], boxes[e][1] + 50), label="可信" if s == "quality" else None)
    poly_arrow(d, [(590, 760), (590, 680), (1690, 680), (1690, 340)], label="不可信则重拍/调整", label_pos=(1260, 680))

    d.text((96, 1190), "设备运行本质：硬件负责形成并采集光斑，软件负责把光斑变化换算成镜片参数。", font=F_TAG, fill=hex_to_rgb(COLORS["muted"]))
    path = save_canvas("03_整个设备运行流程.png")
    img.save(path)
    return path


def set_font_run(run, size=11, bold=False, color="0B2545"):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def build_doc(image_paths):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("焦度计项目三类流程图")
    set_font_run(r, size=22, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("展示整个项目制作流程、软件部分运行流程、整个设备运行流程")
    set_font_run(r, size=12, color="555555")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("用途：项目答辩、团队培训、PPT 截图引用")
    set_font_run(r, size=11, color="666666")

    for i, path in enumerate(image_paths, 1):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(9.9))

    doc.save(DOCX_PATH)
    return DOCX_PATH


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    paths = [project_flow(), software_flow(), device_flow()]
    print(build_doc(paths))
    for p in paths:
        print(p)


if __name__ == "__main__":
    main()
