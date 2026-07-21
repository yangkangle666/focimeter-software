from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


OUT_DIR = Path(r"C:\Users\yangkangle\OneDrive\Desktop\焦度计\outputs")
OUT_PATH = OUT_DIR / "焦度计项目算法与内容答辩手册.docx"


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color, before, after in [
        ("Heading 1", 16, "1F4D78", 14, 7),
        ("Heading 2", 13, "2E74B5", 10, 5),
        ("Heading 3", 11.5, "1F4D78", 7, 3),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("焦度计项目答辩手册")
    set_run_font(r, size=9, color="666666")


def para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = para(doc, text, style=style)
    return p


def number(doc, text, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    p = para(doc, text, style=style)
    return p


def heading(doc, text, level=1):
    return para(doc, text, style=f"Heading {level}")


def callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    tc_pr.append(shd)
    p = cell.paragraphs[0]
    r = p.add_run(title + "：")
    set_run_font(r, bold=True, color="1F4D78")
    r2 = p.add_run(body)
    set_run_font(r2)
    doc.add_paragraph()


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = True
    hdr_cells = t.rows[0].cells
    tr_pr = t.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    for idx, h in enumerate(headers):
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, bold=True, color="1F4D78")
        tc_pr = hdr_cells[idx]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E8EEF5")
        tc_pr.append(shd)
        if widths:
            hdr_cells[idx].width = Inches(widths[idx])
    for row in rows:
        table_row = t.add_row()
        tr_pr = table_row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        cells = table_row.cells
        for idx, val in enumerate(row):
            p = cells[idx].paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, size=9.5)
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()
    return t


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("焦度计项目算法与内容答辩手册")
    set_run_font(r, size=22, bold=True, color="0B2545")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("团队学习版｜用于项目方向答辩、算法答辩和后续接手准备")
    set_run_font(r2, size=12, color="555555")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("适用对象：项目负责人、算法组、软件组、汇报成员")
    set_run_font(r3, size=10, color="666666")
    doc.add_paragraph()
    callout(
        doc,
        "本手册的目标",
        "让团队成员在答辩前形成统一口径：知道项目是什么、为什么这样做、软件最后要交付什么、当前原型还缺什么、老师可能怎么问以及我们怎么答。"
    )


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_page_number(doc.sections[0])
    add_title(doc)

    heading(doc, "0. 先把一句话说清楚", 1)
    callout(
        doc,
        "项目一句话",
        "本项目要完成自动焦度计的软件算法部分：利用硬件采集到的标定图像和测量图像，识别哈特曼光斑的位置变化，计算并输出镜片的球镜度 S、柱镜度 C 和轴位 A。"
    )
    para(doc, "如果老师问“你们到底做什么”，不要先讲一堆图像处理名词。先说：我们做的是自动焦度计的软件算法，让硬件拍到的光斑图片变成可读的屈光度结果。")
    table(doc, ["容易混淆的问题", "统一回答"], [
        ["是不是做整台仪器？", "不是一开始就做整台仪器。项目整体是自动焦度计，包含硬件和软件；我们团队主要承担软件算法部分。"],
        ["软件输入是什么？", "至少包括标定图像、测量图像、硬件参数或配置参数。后续可扩展到相机实时采集。"],
        ["软件输出是什么？", "核心输出是 S、C、A，即球镜度、柱镜度、轴位；同时应输出识别质量、异常提示和中间过程图。"],
        ["当前代码是什么状态？", "当前是算法流程原型，能够演示从两张图片到度数计算的主链路，但还需要工程化、参数化和验证。"],
    ], widths=[1.8, 4.7])

    heading(doc, "1. 答辩目标与评审关心点", 1)
    para(doc, "这次答辩的目的不是证明我们已经做完产品，而是证明我们理解项目、路线可行、团队能接手，并且知道后续如何把原型变成可靠的软件。")
    table(doc, ["评审关心点", "我们要证明什么"], [
        ["项目理解", "我们知道自动焦度计是什么，也知道本团队边界是软件算法。"],
        ["技术路线", "图像处理、光斑质心、坐标转换、屈光度计算这条链路是连贯的。"],
        ["可行性", "已有 focimeter 原型可作为流程依据，但我们知道它仍是原型。"],
        ["风险意识", "我们不回避标定、误差、稳定性、参数来源、标准验证这些难点。"],
        ["团队组织", "8 人团队有明确分工，能从资料理解推进到可执行软件。"],
    ], widths=[1.8, 4.7])

    heading(doc, "2. 项目背景：自动焦度计是什么", 1)
    para(doc, "焦度计用于测量眼镜镜片的屈光参数。人工或半自动设备依赖操作者读取和调节，自动焦度计则希望把光学成像、图像识别和参数计算结合起来，让镜片参数自动输出。")
    bullet(doc, "硬件部分：光源、准直光路、镜片支架、哈特曼光阑或类似采样结构、成像镜头、相机。")
    bullet(doc, "软件部分：读取图像，识别光斑，比较标定图和测量图的位移，计算屈光参数，输出结果并判断结果是否可信。")
    bullet(doc, "本团队的重点：软件算法。硬件参数会影响计算，所以软件必须把硬件参数配置化，而不是把常数写死。")

    heading(doc, "3. 最终软件应该做成什么", 1)
    callout(
        doc,
        "可运行算法软件的定义",
        "不是只写一段公式，也不是只做一个界面。它应该能接收标定图、测量图和配置参数，自动识别光斑，计算 S/C/A，输出结果、日志、中间图和异常提示。"
    )
    table(doc, ["模块", "应该具备的能力", "当前原型状态"], [
        ["输入模块", "导入图片或接入相机；支持标定图和测量图；支持配置文件。", "当前路径写死，适合演示，不适合交付。"],
        ["图像处理模块", "ROI、去噪、增强、二值化、连通域、质心定位。", "已有基本流程，但参数和鲁棒性要增强。"],
        ["标定模块", "建立坐标系，保存标定结果，支持重复调用。", "已有坐标系建立，但保存/加载不足。"],
        ["计算模块", "根据光斑位移和光路参数计算 S/C/A。", "已有球镜/柱镜计算雏形，参数和公式依据需要整理。"],
        ["结果模块", "输出结果、质量评分、中间图、错误原因、测试报告。", "目前主要是控制台输出。"],
        ["验证模块", "标准镜片测试、重复性、误差统计、与标准设备对比。", "还需要系统补齐。"],
    ], widths=[1.45, 2.65, 2.4])

    heading(doc, "4. 当前 focimeter 原型在做什么", 1)
    para(doc, "可以把 focimeter.cpp 理解为算法总控原型。它串起了标定、测量、坐标转换和屈光度计算，适合作为答辩时解释“我们软件链路怎么跑”的依据。")
    number(doc, "读取标定图像和测量图像。")
    number(doc, "分别处理图像，提取 5 个光斑连通域。")
    number(doc, "根据标定图中的 5 个点建立局部坐标系。")
    number(doc, "把标定图和测量图中的关键光斑转到同一坐标系。")
    number(doc, "比较测量图光斑相对标定图的位移。")
    number(doc, "判断球镜或柱镜，并调用对应计算函数输出结果。")
    callout(
        doc,
        "答辩说法",
        "当前原型不是完整产品，而是算法链路验证样机。它证明路线可走，但后续必须补配置、数据、误差验证、异常处理和结果输出。"
    )

    heading(doc, "5. 核心算法路线", 1)
    para(doc, "整个软件算法可以按下面这条链路理解：")
    table(doc, ["步骤", "输入", "处理", "输出"], [
        ["1 图像获取", "标定图、测量图", "读取图片或相机帧", "原始图像"],
        ["2 图像预处理", "原始图像", "ROI、灰度化、中值滤波、顶帽增强", "增强后的光斑图"],
        ["3 光斑分割", "增强图像", "区域 Otsu、连通域筛选", "5 个候选光斑"],
        ["4 质心定位", "光斑区域", "灰度加权矩或轮廓矩", "每个光斑的亚像素/像素质心"],
        ["5 坐标标定", "标定图 5 点", "中心点设原点，方向点定义 X/Y 轴", "标定坐标系"],
        ["6 位移计算", "标定点与测量点", "统一坐标系下做差", "光斑偏移量"],
        ["7 度数计算", "偏移量、像元尺寸、光路距离", "代入光学模型", "S/C/A"],
        ["8 质量判断", "识别数量、残差、重复性", "阈值与异常检测", "可信度与错误提示"],
    ], widths=[1.1, 1.45, 2.55, 1.4])

    heading(doc, "6. 为什么用光斑位移来算镜片度数", 1)
    para(doc, "镜片会改变光线传播方向。没有镜片时，光斑落在参考位置；放入镜片后，光束发生偏折，光斑位置随之变化。只要光路结构和相机像元尺寸已知，就可以把图像上的位移换算成光线偏折，再进一步换算成镜片屈光度。")
    bullet(doc, "标定图的作用：提供无镜片或参考状态下的光斑位置，也就是“零点”或基准。")
    bullet(doc, "测量图的作用：提供放入镜片后的光斑位置。")
    bullet(doc, "两者差异：反映镜片造成的光线偏折。")
    bullet(doc, "软件计算：把像素位移转成物理位移，再结合光路距离计算屈光度。")
    callout(
        doc,
        "注意",
        "陈文婷文章可以用于解释“为什么哈特曼光阑、图像预处理、质心识别、标定比较是合理路线”，但本项目的计算公式、参数和结果不要直接照搬她的文章。"
    )

    heading(doc, "7. 图像处理环节怎么讲", 1)
    table(doc, ["环节", "作用", "答辩解释"], [
        ["ROI", "减少无关背景", "只处理中心有效区域，提高速度并降低噪声干扰。"],
        ["灰度化", "简化数据", "光斑识别主要依赖亮度信息，不需要彩色通道。"],
        ["中值滤波", "抑制椒盐噪声", "保留边缘的同时减少孤立噪点。"],
        ["顶帽运算", "增强亮光斑", "突出比背景更亮的小目标，适合光斑提取。"],
        ["区域 Otsu", "自适应分割", "不同区域亮度可能不均匀，局部分割比全局阈值更稳。"],
        ["连通域", "找到光斑区域", "筛选主要连通区域，保留目标光斑。"],
        ["质心", "确定光斑位置", "用光斑灰度分布的中心代表光束落点。"],
    ], widths=[1.25, 1.55, 3.7])
    para(doc, "如果老师问为什么不是深度学习：本项目目标是精密测量，样本量可能有限，且需要可解释的物理量。传统图像处理加光学模型更容易解释、标定和追溯误差。后续如果数据充足，可以把深度学习用于光斑检测辅助，而不是一开始替代物理模型。")

    heading(doc, "8. 坐标系和标定怎么讲", 1)
    para(doc, "标定不是可有可无的步骤。相机可能有旋转，光斑阵列不一定和图像水平垂直完全一致，所以要在标定图里建立一个局部坐标系。")
    bullet(doc, "中心光斑作为原点。")
    bullet(doc, "某个方向光斑作为 Y 轴方向。")
    bullet(doc, "另一个方向光斑作为 X 轴方向。")
    bullet(doc, "测量图中的光斑统一转换到这个坐标系下比较，避免相机安装角度影响结果。")
    callout(
        doc,
        "答辩重点",
        "标定的本质是把“相机像素坐标”变成“可比较的局部测量坐标”。没有标定，测量图和参考图的位移很容易混入相机安装偏差。"
    )

    heading(doc, "9. 球镜、柱镜、轴位怎么讲", 1)
    para(doc, "镜片参数不要一上来就讲复杂公式。先把物理意义讲清楚。")
    table(doc, ["参数", "含义", "软件如何得到"], [
        ["球镜度 S", "各方向基本一致的会聚或发散能力。", "主要看光斑沿主方向的整体位移变化。"],
        ["柱镜度 C", "不同方向屈光力不同，存在散光。", "比较两个方向的位移差异，得到柱镜分量。"],
        ["轴位 A", "柱镜作用的方向角。", "由两个方向的位移关系计算角度。"],
    ], widths=[1.2, 2.3, 3.0])
    para(doc, "当前代码里有球镜计算 Slens 和柱镜计算 Clens，但这里要谨慎表达：原型给出了计算形式，后续接手后需要根据实际光路参数、标准镜片和检定标准重新校核参数和公式。")

    heading(doc, "10. 陈文婷文章该怎么用", 1)
    table(doc, ["可以参考", "不要参考"], [
        ["哈特曼法/光斑位移法的背景解释。", "不要直接引用她的具体计算参数作为本项目参数。"],
        ["图像预处理、光斑分割、质心定位的思路。", "不要把她的实验条件默认等同于我们的硬件条件。"],
        ["为什么这种路线适合镜片屈光力测量。", "不要把她的实验结果当成本项目结果。"],
        ["论文中的方法局限可作为风险意识。", "不要说我们计算完全照搬某篇文章。"],
    ], widths=[3.25, 3.25])
    callout(
        doc,
        "统一口径",
        "陈文婷文章用于方法背景和思路理解；本项目计算依据应来自本项目硬件参数、光学模型、标定实验和相关标准。"
    )

    heading(doc, "11. 当前原型还需要完善什么", 1)
    table(doc, ["缺口", "为什么重要", "后续怎么补"], [
        ["输入不工程化", "路径写死，无法适应真实使用。", "改成命令行参数、配置文件、批量导入或相机接口。"],
        ["参数写死", "像元尺寸、距离、阈值会随硬件变化。", "建立配置文件和参数校准流程。"],
        ["识别鲁棒性不足", "曝光、噪声、遮挡会导致找不到 5 个点。", "增加质量评分、异常检测、候选点排序和失败提示。"],
        ["公式未闭环验证", "计算结果需要和真实镜片对应。", "用标准镜片建立误差表，修正模型参数。"],
        ["结果输出简单", "控制台输出不利于交付和复盘。", "输出报告、CSV/JSON、中间图、日志。"],
        ["没有测试集", "无法证明稳定性和精度。", "建立样本库，覆盖球镜、柱镜、不同度数和异常样本。"],
        ["没有用户界面", "非开发人员难以使用。", "先做离线算法工具，再考虑 GUI。"],
    ], widths=[1.35, 2.2, 2.95])

    heading(doc, "12. 8 人团队建议分工", 1)
    table(doc, ["角色", "人数", "主要任务"], [
        ["项目负责人", "1", "统一口径、答辩主线、进度管理、风险和交付边界控制。"],
        ["算法原理组", "2", "整理光学模型、公式变量、标定关系、标准镜片验证方案。"],
        ["图像处理组", "2", "负责 ROI、滤波、二值化、连通域、质心定位和鲁棒性优化。"],
        ["软件工程组", "2", "把原型改成可运行程序，负责配置、输入输出、日志和测试框架。"],
        ["资料与答辩组", "1", "维护 PPT、答辩手册、Q&A、演练记录和资料归档。"],
    ], widths=[1.5, 0.8, 4.2])
    para(doc, "负责人答辩时要表现为“能整合团队”，不是每个公式都自己讲到最深。你要掌握边界、路线、风险和下一步。细节问题可以指定算法组或图像组补充。")

    heading(doc, "13. 推荐答辩汇报结构", 1)
    table(doc, ["页码", "页面标题", "要讲出的结论"], [
        ["1", "我们要做自动焦度计的软件算法", "项目不是泛泛做图像处理，而是把光斑图片转成镜片屈光参数。"],
        ["2", "硬件拍图，软件算数", "硬件负责形成并采集光斑，软件负责识别位移并计算 S/C/A。"],
        ["3", "核心物理量是光斑位移", "镜片改变光线方向，导致测量图光斑相对标定图移动。"],
        ["4", "算法链路已经形成原型", "focimeter 原型串起了图像处理、坐标标定和参数计算。"],
        ["5", "图像处理保证光斑找得准", "ROI、滤波、顶帽、二值化、连通域、质心定位。"],
        ["6", "标定保证坐标可比较", "用标定图建立局部坐标系，消除相机安装角度影响。"],
        ["7", "计算输出 S/C/A", "把位移和光路参数结合，得到球镜度、柱镜度、轴位。"],
        ["8", "当前原型和待完善内容", "我们知道缺口：配置、鲁棒性、验证、输出。"],
        ["9", "团队分工和接手计划", "8 人分工明确，先离线算法，再工程化，再验证。"],
        ["10", "结论", "我们能接手软件部分，并有清晰路线把原型变成可运行算法软件。"],
    ], widths=[0.7, 2.1, 3.7])

    heading(doc, "14. 开场 1 分钟模板", 1)
    para(doc, "各位老师好，我们团队这次申请接手的是自动焦度计项目的软件算法部分。自动焦度计的核心目标，是让硬件采集到的光斑图像，经过软件处理后自动输出镜片的球镜度、柱镜度和轴位。我们目前已经梳理了项目资料和现有 focimeter 原型，明确了软件链路：先用标定图建立参考坐标系，再处理测量图，识别光斑质心，计算光斑位移，最后结合光路参数输出屈光结果。")
    para(doc, "我们也清楚当前原型还不是完整产品，后续需要补齐参数配置、稳定识别、标定保存、误差验证和结果输出。今天的答辩重点不是说项目已经完成，而是说明我们理解项目、路线可行、风险明确，并且团队有能力把现有原型推进成可运行的算法软件。")

    heading(doc, "15. 老师可能问的问题与建议回答", 1)
    qa_rows = [
        ["项目边界", "你们到底做硬件还是软件？", "项目整体包含硬件和软件，但我们团队主要负责软件算法部分。硬件提供标定图和测量图，软件负责识别光斑位移并计算镜片参数。"],
        ["项目边界", "最后交付物是什么？", "阶段性交付是可运行算法软件：输入标定图、测量图和配置参数，输出 S/C/A、识别质量、中间图和日志。"],
        ["项目边界", "这和普通图像识别有什么区别？", "普通图像识别只关心目标位置，本项目要把位置变化转成有物理意义的屈光度，所以必须结合光路参数和标定。"],
        ["原理", "为什么镜片会让光斑移动？", "镜片改变光线传播方向，经过固定距离后，光束落点发生变化。光斑位移反映光线偏折，进而反映屈光力。"],
        ["原理", "为什么需要标定图？", "标定图提供参考位置和坐标系。没有标定，测量图里的位移会混入相机安装角度、光斑阵列偏差等因素。"],
        ["原理", "标定图一定是无镜片图吗？", "通常可理解为无镜片或标准参考状态图。具体要根据硬件实验方案确认，但它的作用是建立参考基准。"],
        ["算法", "为什么要找 5 个点？", "当前原型用中心点和四周方向点建立局部坐标系并提取两个主方向的位移。后续可根据实际光阑结构扩展点数。"],
        ["算法", "为什么用质心？", "光斑不是一个单像素点，质心能综合光斑亮度分布，作为光束落点估计，比取最大亮点更稳定。"],
        ["算法", "为什么用顶帽运算？", "顶帽适合增强比背景更亮的小目标，能突出光斑并减弱不均匀背景。"],
        ["算法", "为什么用 Otsu？", "Otsu 能自动寻找阈值。区域 Otsu 可以适应局部亮度不均匀，比一个全局阈值更稳。"],
        ["算法", "如果找不到 5 个光斑怎么办？", "当前原型会失败退出。后续软件要增加异常提示、曝光检查、候选点补救和质量评分。"],
        ["算法", "球镜和柱镜怎么区分？", "球镜各方向偏折较一致；柱镜不同方向偏折不同。软件通过关键方向上的位移关系进行初步分类。"],
        ["算法", "轴位怎么得到？", "轴位来自柱镜两个主方向位移之间的角度关系。实际实现需要结合光路模型和标定参数校核。"],
        ["公式", "你们的公式依据是什么？", "公式应来自光学成像关系、本项目光路参数和标定实验。论文只作为方法背景，不能直接替代本项目参数。"],
        ["公式", "陈文婷论文是不是你们计算依据？", "不是。我们参考她对哈特曼法、图像处理和质心识别的思路说明，但计算参数和结果必须基于本项目硬件和实验校准。"],
        ["公式", "代码里的像元尺寸和距离可靠吗？", "当前是原型硬编码参数，用于流程验证。接手后会改为配置参数，并通过硬件规格和标定实验确认。"],
        ["工程", "当前软件能直接交付吗？", "不能直接交付。它是算法原型，需要补输入配置、鲁棒性、验证、报告和用户操作方式。"],
        ["工程", "为什么先做离线算法软件？", "离线软件能先验证算法正确性，不受相机实时采集影响。算法稳定后再接相机和界面，风险更低。"],
        ["工程", "你们如何保证可维护？", "把图像处理、标定、计算、输出分模块；参数配置化；建立测试集和日志，让问题可复现。"],
        ["验证", "怎么证明测得准？", "用已知度数的标准镜片或标准焦度计结果对比，统计误差、重复性和失败率。"],
        ["验证", "要满足 JJG 580 吗？", "如果目标是计量级产品，需要对照相关检定规程；如果阶段目标是研究原型，先完成算法可行性验证，再逐步靠近标准要求。"],
        ["风险", "最大风险是什么？", "最大风险是硬件条件和参数不明确导致公式无法闭环，以及图像质量变化导致光斑识别不稳定。"],
        ["风险", "如何降低风险？", "先确认硬件参数和标定图定义；建立样本库；把参数配置化；用标准镜片做误差表。"],
        ["团队", "8 个人怎么分工？", "负责人统筹，算法原理组管公式和标定，图像处理组管识别，软件工程组管可运行程序，资料答辩组管文档和演练。"],
        ["团队", "如果通过答辩，第一步做什么？", "第一步不是盲目改代码，而是确认硬件参数、样本数据、标定图定义和最终验收指标。"],
    ]
    table(doc, ["类别", "可能问题", "建议回答"], qa_rows, widths=[0.95, 2.0, 3.55])

    heading(doc, "16. 答辩时不要踩的坑", 1)
    table(doc, ["不要这样说", "建议这样说"], [
        ["我们已经做完了。", "我们已有算法流程原型，后续还要工程化和验证。"],
        ["公式就是论文里的。", "论文用于方法背景，计算要基于本项目光路和标定实验。"],
        ["只要识别出光斑就行。", "识别只是第一步，还要标定、计算、误差验证和结果质量判断。"],
        ["硬件不归我们管，所以不用管参数。", "软件不设计硬件，但必须读取和使用硬件参数，否则无法准确计算。"],
        ["深度学习更高级，所以后续用深度学习。", "本项目优先采用可解释的图像处理和光学模型，深度学习可作为后续辅助。"],
        ["老师问精度时直接承诺很高。", "先说明需要标准镜片和检定数据验证，再给阶段目标和验证方案。"],
    ], widths=[3.0, 3.5])

    heading(doc, "17. 接手后的路线图", 1)
    table(doc, ["阶段", "目标", "关键产出"], [
        ["第 1 阶段：澄清与复现", "确认硬件参数、样本图、标定图定义，复现当前原型。", "项目规格说明、样本清单、可复现运行记录。"],
        ["第 2 阶段：离线算法工具", "把写死路径和参数改成配置，稳定处理图片。", "命令行程序、配置文件、结果 CSV/JSON、中间图输出。"],
        ["第 3 阶段：算法校准", "用标准镜片校核公式和参数，建立误差统计。", "误差报告、参数修正方案、异常样本分析。"],
        ["第 4 阶段：软件化交付", "补日志、界面或批处理、测试集和使用说明。", "可运行软件包、用户说明、测试报告。"],
        ["第 5 阶段：硬件联调", "接入相机和真实采集流程。", "端到端演示：采图到输出结果。"],
    ], widths=[1.45, 2.45, 2.6])

    heading(doc, "18. 团队学习清单", 1)
    para(doc, "答辩前每个成员至少要掌握下面这些内容。负责人尤其要能把所有内容串起来。")
    bullet(doc, "能用一句话说清楚项目：硬件拍光斑，软件算 S/C/A。")
    bullet(doc, "知道标定图和测量图分别有什么作用。")
    bullet(doc, "知道软件算法链路：预处理、分割、质心、坐标、位移、计算、验证。")
    bullet(doc, "知道当前原型只是流程验证，不是完整交付。")
    bullet(doc, "知道陈文婷文章能参考什么，不能参考什么。")
    bullet(doc, "能回答“最后做出什么软件”。")
    bullet(doc, "能说明后续为什么要做标准镜片验证。")
    bullet(doc, "能说出自己负责哪一块，以及和其他组的接口是什么。")

    heading(doc, "19. 负责人压轴总结模板", 1)
    para(doc, "我们团队对这个项目的理解是：自动焦度计不是单纯的软件项目，也不是单纯的硬件项目，而是光学成像、图像处理和参数计算结合的系统。我们本阶段主要承担软件算法部分，目标是把硬件采集到的标定图和测量图转化为镜片的 S/C/A 结果。")
    para(doc, "目前已有 focimeter 原型可以说明算法链路基本成立，但它还停留在原型阶段，存在路径写死、参数硬编码、识别鲁棒性和验证不足等问题。我们通过答辩后，会先确认硬件参数和样本数据，再把原型改造成可配置、可测试、可输出报告的离线算法软件，最后结合标准镜片和真实采集逐步完成验证。")
    para(doc, "因此，我们不是只会讲概念，而是已经明确了项目边界、算法路线、风险点、团队分工和接手后的执行路径。")

    heading(doc, "20. 术语速查", 1)
    table(doc, ["术语", "解释"], [
        ["焦度计", "测量眼镜镜片屈光参数的仪器。"],
        ["自动焦度计", "通过光学成像和软件算法自动输出镜片参数的焦度计。"],
        ["球镜度 S", "镜片整体会聚或发散能力，单位通常为 D。"],
        ["柱镜度 C", "散光相关参数，表示不同方向屈光力差异。"],
        ["轴位 A", "柱镜作用方向的角度。"],
        ["标定图", "参考状态下的光斑图，用于建立基准。"],
        ["测量图", "放入待测镜片后的光斑图。"],
        ["质心", "光斑亮度分布的中心，用来代表光束落点。"],
        ["ROI", "感兴趣区域，只处理图像中有效部分。"],
        ["顶帽运算", "突出亮小目标、抑制不均匀背景的形态学处理。"],
        ["连通域", "二值图中互相连接的一片目标区域。"],
        ["鲁棒性", "算法在噪声、曝光变化、位置偏差下仍能稳定工作的能力。"],
    ], widths=[1.6, 4.9])

    doc.core_properties.title = "焦度计项目算法与内容答辩手册"
    doc.core_properties.subject = "自动焦度计软件算法答辩资料"
    doc.core_properties.author = "Codex"
    doc.save(OUT_PATH)
    print(str(OUT_PATH))


if __name__ == "__main__":
    build()
