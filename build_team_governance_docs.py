from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from build_defense_manual import callout, heading, table, set_run_font


ROOT = Path(r"C:\Users\yangkangle\OneDrive\Desktop\焦度计")
OUT_DIR = ROOT / "outputs"

HUMAN_MD = OUT_DIR / "焦度计项目团队总原则_人看版.md"
HUMAN_DOCX = OUT_DIR / "焦度计项目团队总原则_人看版.docx"
AI_MD = OUT_DIR / "焦度计项目AI开发总规范_AI看版.md"
AI_DOCX = OUT_DIR / "焦度计项目AI开发总规范_AI看版.docx"


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    for name, size, color, before, after in [
        ("Heading 1", 16, "1F4D78", 14, 6),
        ("Heading 2", 12.5, "2E74B5", 10, 4),
        ("Heading 3", 11.5, "1F4D78", 7, 3),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, name="Microsoft YaHei", size=22, bold=True, color="0B2545")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    set_run_font(r2, name="Microsoft YaHei", size=11.5, color="666666")
    doc.add_paragraph()


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, name="Microsoft YaHei", size=10.5)


def add_numbers(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, name="Microsoft YaHei", size=10.5)


def add_para(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, name="Microsoft YaHei", size=10.5)


def md_table(headers, rows):
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        out.append("| " + " | ".join(row) + " |")
    return "\n".join(out)


def write_md(path: Path, lines):
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_human_md():
    lines = []
    lines.append("# 焦度计项目团队总原则（人看版）")
    lines.append("")
    lines.append("> 这是给团队所有成员看的统一规则。它只讲三件事：总原则、四大独立板块、三个开发阶段。")
    lines.append("")
    lines.append("## 一句话先记住")
    lines.append("")
    lines.append("四个板块必须独立开发，统一接口，先做本地系统，再合并成整体，最后再优化成软件。任何人都不能擅自改接口、改单位、改坐标、改输出格式。")
    lines.append("")
    lines.append("## 总原则")
    lines.append("")
    lines.append("- 四个板块各自独立，互不影响进度。")
    lines.append("- 所有合并都靠统一接口，不靠临时手工拼接。")
    lines.append("- 第一阶段先做本地系统，不追求最终软件外壳。")
    lines.append("- 第二阶段合并并修 bug。")
    lines.append("- 第三阶段再做优化创新，把本地系统逐步变成软件。")
    lines.append("- 任何统一项都不允许任何人私自改动。")
    lines.append("")
    lines.append("## 四大独立板块")
    lines.append("")
    lines.append(md_table(
        ["板块", "主要功能", "和其他板块的关系", "第一阶段独立产出"],
        [
            ["M1 输入与配置", "管理标定图、测量图、配置文件、样本路径", "给 M2/M3/M4 提供标准输入", "标准输入包和配置读取结果"],
            ["M2 图像识别", "预处理图片、找光斑、算质心", "接收 M1 数据，输出给 M3", "标准光斑坐标 JSON"],
            ["M3 标定与计算", "建坐标系、算位移、算 S/C/A", "接收 M2 坐标和 M1 参数", "镜片结果 JSON"],
            ["M4 本地系统与展示", "主流程、展示、日志、导出、异常提示", "串联前三块，形成完整本地系统", "可演示的本地系统外壳"],
        ],
    ))
    lines.append("")
    lines.append("## 三个开发阶段")
    lines.append("")
    lines.append("### 第一阶段：四个板块同时开工")
    lines.append("")
    lines.append("- 四组同步开发，彼此不等。")
    lines.append("- 每组都要先做出自己的模块成品。")
    lines.append("- 允许使用 mock 数据、样例数据、模拟接口。")
    lines.append("- 验收重点是“能独立运行”，不是“已经集成完美”。")
    lines.append("")
    lines.append("### 第二阶段：合并并修复 bug")
    lines.append("")
    lines.append("- 按统一接口把四个模块合成完整本地系统。")
    lines.append("- 优先修接口问题、字段问题、单位问题、坐标问题。")
    lines.append("- 这一阶段不优先加新功能。")
    lines.append("")
    lines.append("### 第三阶段：优化创新")
    lines.append("")
    lines.append("- 在完整本地系统基础上做界面、封装、导出、自动化。")
    lines.append("- 后续把本地系统升级成正式软件。")
    lines.append("- 不能推翻前两阶段已经统一好的接口。")
    lines.append("")
    lines.append("## 一定要统一的点")
    lines.append("")
    lines.append(md_table(
        ["统一项", "为什么必须统一", "任何人都不能擅改什么"],
        [
            ["输入输出格式", "不统一就无法合并", "字段名、层级、文件格式"],
            ["单位", "不统一就会算错", "像素、毫米、米、角度的写法"],
            ["坐标系", "不统一就会方向错", "原点、X/Y 正方向、转换方式"],
            ["错误格式", "不统一就无法排查", "错误码、错误消息、返回结构"],
            ["日志格式", "不统一就无法复现", "文件名、时间、参数、结果"],
            ["命名规则", "不统一就会乱", "文件名、变量名、结果字段名"],
            ["样本集", "不统一就无法对比", "测试图片、异常图片、标准镜片样本"],
            ["版本号", "不统一就无法追踪", "模块版本、变更记录"],
            ["目录结构", "不统一就难合并", "输入、输出、日志、配置目录"],
        ],
    ))
    lines.append("")
    lines.append("## 绝不可以擅自改动的事情")
    lines.append("")
    lines.append("- 不准私自改接口。")
    lines.append("- 不准私自改字段名。")
    lines.append("- 不准私自改单位。")
    lines.append("- 不准私自改坐标系。")
    lines.append("- 不准私自改输出格式。")
    lines.append("- 不准把个人电脑路径写死进正式代码。")
    lines.append("- 不准等别人模块完成后才开始自己的模块。")
    lines.append("- 不准在没有说明的情况下偷偷换实现方式。")
    lines.append("")
    lines.append("## 最后记一句")
    lines.append("")
    lines.append("四个板块可以分别开发，但必须先统一接口和约束。先各自做成品，再合并成系统，最后再优化成软件。")
    return lines


def build_human_docx():
    doc = Document()
    setup_doc(doc)
    add_title(doc, "焦度计项目团队总原则（人看版）", "给团队所有成员看的统一规则：先独立、再合并、后优化")
    callout(
        doc,
        "先记住这句话",
        "四个板块必须独立开发、统一接口、先做本地系统，再合并成整体，最后再优化成软件。任何人都不能擅自改接口、改单位、改坐标、改输出格式。"
    )

    heading(doc, "总原则", 1)
    add_bullets(doc, [
        "四个板块各自独立，互不影响进度。",
        "所有合并都靠统一接口，不靠临时手工拼接。",
        "第一阶段先做本地系统，不追求最终软件外壳。",
        "第二阶段合并并修 bug。",
        "第三阶段再做优化创新，把本地系统逐步变成软件。",
        "任何统一项都不允许任何人私自改动。",
    ])

    heading(doc, "四大独立板块", 1)
    table(doc,
          ["板块", "主要功能", "和其他板块的关系", "第一阶段独立产出"],
          [
              ["M1 输入与配置", "管理标定图、测量图、配置文件、样本路径", "给 M2/M3/M4 提供标准输入", "标准输入包和配置读取结果"],
              ["M2 图像识别", "预处理图片、找光斑、算质心", "接收 M1 数据，输出给 M3", "标准光斑坐标 JSON"],
              ["M3 标定与计算", "建坐标系、算位移、算 S/C/A", "接收 M2 坐标和 M1 参数", "镜片结果 JSON"],
              ["M4 本地系统与展示", "主流程、展示、日志、导出、异常提示", "串联前三块，形成完整本地系统", "可演示的本地系统外壳"],
          ],
          widths=[1.5, 2.25, 1.9, 1.85])

    heading(doc, "三个开发阶段", 1)
    heading(doc, "第一阶段：四个板块同时开工", 2)
    add_bullets(doc, [
        "四组同步开发，彼此不等。",
        "每组都要先做出自己的模块成品。",
        "允许使用 mock 数据、样例数据、模拟接口。",
        "验收重点是“能独立运行”，不是“已经集成完美”。",
    ])
    heading(doc, "第二阶段：合并并修复 bug", 2)
    add_bullets(doc, [
        "按统一接口把四个模块合成完整本地系统。",
        "优先修接口问题、字段问题、单位问题、坐标问题。",
        "这一阶段不优先加新功能。",
    ])
    heading(doc, "第三阶段：优化创新", 2)
    add_bullets(doc, [
        "在完整本地系统基础上做界面、封装、导出、自动化。",
        "后续把本地系统升级成正式软件。",
        "不能推翻前两阶段已经统一好的接口。",
    ])

    heading(doc, "一定要统一的点", 1)
    table(doc,
          ["统一项", "为什么必须统一", "任何人都不能擅改什么"],
          [
              ["输入输出格式", "不统一就无法合并", "字段名、层级、文件格式"],
              ["单位", "不统一就会算错", "像素、毫米、米、角度的写法"],
              ["坐标系", "不统一就会方向错", "原点、X/Y 正方向、转换方式"],
              ["错误格式", "不统一就无法排查", "错误码、错误消息、返回结构"],
              ["日志格式", "不统一就无法复现", "文件名、时间、参数、结果"],
              ["命名规则", "不统一就会乱", "文件名、变量名、结果字段名"],
              ["样本集", "不统一就无法对比", "测试图片、异常图片、标准镜片样本"],
              ["版本号", "不统一就无法追踪", "模块版本、变更记录"],
              ["目录结构", "不统一就难合并", "输入、输出、日志、配置目录"],
          ],
          widths=[1.25, 2.15, 2.8])

    heading(doc, "绝不可以擅自改动的事情", 1)
    add_bullets(doc, [
        "不准私自改接口。",
        "不准私自改字段名。",
        "不准私自改单位。",
        "不准私自改坐标系。",
        "不准私自改输出格式。",
        "不准把个人电脑路径写死进正式代码。",
        "不准等别人模块完成后才开始自己的模块。",
        "不准在没有说明的情况下偷偷换实现方式。",
    ])

    heading(doc, "最后记一句", 1)
    add_para(doc, "四个板块可以分别开发，但必须先统一接口和约束。先各自做成品，再合并成系统，最后再优化成软件。")
    doc.save(HUMAN_DOCX)


def build_ai_md():
    lines = []
    lines.append("# 焦度计项目 AI 开发总规范（AI 看版）")
    lines.append("")
    lines.append("> 这份文档是给 AI 开工前读的。目标不是讲概念，而是约束 AI 怎样开发、怎样统一接口、怎样为后续合并服务。")
    lines.append("")
    lines.append("## 开工前必须先问成员的固定问题")
    lines.append("")
    lines.append("AI 在开始写代码前，必须先问成员下面这 4 个问题：")
    lines.append("")
    lines.append("1. 你负责哪个板块？M1 输入与配置 / M2 图像识别 / M3 标定与计算 / M4 本地系统与展示。")
    lines.append("2. 你现在处于哪个阶段？第一阶段模块成品 / 第二阶段合并修 bug / 第三阶段优化创新。")
    lines.append("3. 你希望我先输出什么？代码 / 接口 / 测试样例 / 文档 / 调试建议。")
    lines.append("4. 你手里现在有哪些本地样本、配置和 mock 数据？")
    lines.append("")
    lines.append("如果成员没有回答清楚，AI 不能擅自假设，必须继续追问。")
    lines.append("")
    lines.append("## 项目目标")
    lines.append("")
    lines.append("项目目标是把自动焦度计的图片输入变成可读的镜片参数输出。")
    lines.append("")
    lines.append("```text")
    lines.append("输入：标定图 + 测量图 + 配置参数")
    lines.append("处理：识别光斑，建立坐标系，计算位移，输出 S/C/A")
    lines.append("输出：镜片结果、质量状态、日志、中间图、错误码")
    lines.append("```")
    lines.append("")
    lines.append("## 四个独立模块")
    lines.append("")
    lines.append(md_table(
        ["模块", "职责", "输入", "输出", "禁止做什么"],
        [
            ["M1 输入与配置", "管理标定图、测量图、配置文件、样本路径", "图片路径、配置文件、任务编号", "标准输入包 JSON", "不要做图像识别和度数计算"],
            ["M2 图像识别", "预处理、找光斑、算质心", "图片和图像处理参数", "spots JSON", "不要计算 S/C/A"],
            ["M3 标定与计算", "建坐标系、算位移、算 S/C/A", "光斑坐标和光学参数", "result JSON", "不要读取原始图片做识别"],
            ["M4 本地系统与展示", "主流程、展示、日志、导出、异常提示", "各模块输出或 mock 数据", "本地系统成品", "不要重写核心算法"],
        ],
    ))
    lines.append("")
    lines.append("## 必须统一的点")
    lines.append("")
    lines.append(md_table(
        ["统一项", "统一要求"],
        [
            ["输入输出格式", "所有模块都必须使用统一 JSON 结构，字段名不能私改。"],
            ["单位", "像素、毫米、米、角度必须写清楚，不能混用。"],
            ["坐标系", "必须说明原点、X/Y 正方向、图像坐标和测量坐标的关系。"],
            ["错误格式", "错误必须返回统一 error JSON，不能只打印控制台。"],
            ["日志格式", "要保存 task_id、模块名、参数、结果、错误。"],
            ["样本集", "所有组使用同一批样本和 mock 数据进行开发和测试。"],
            ["版本号", "每个模块都要有版本号和变更记录。"],
            ["目录结构", "输入、输出、日志、配置目录要统一命名。"],
            ["运行入口", "每个模块都要能本地独立启动，不能靠手改源码。"],
        ],
    ))
    lines.append("")
    lines.append("## 三个开发阶段")
    lines.append("")
    lines.append("### 第一阶段：四个模块同时开工")
    lines.append("")
    lines.append("- 每个模块都必须先做出自己的可运行成品。")
    lines.append("- 允许使用 mock 数据。")
    lines.append("- 允许先做最小可运行版本。")
    lines.append("- 不能等待别的模块完成。")
    lines.append("")
    lines.append("### 第二阶段：合并并修复 bug")
    lines.append("")
    lines.append("- 通过统一接口把四个模块接成完整本地系统。")
    lines.append("- 优先修接口、字段、单位、坐标和错误处理。")
    lines.append("- 不要在这个阶段大范围重构。")
    lines.append("")
    lines.append("### 第三阶段：优化创新")
    lines.append("")
    lines.append("- 在本地系统稳定后做界面、封装、自动导出、自动测试。")
    lines.append("- 后续把本地系统升级成正式软件。")
    lines.append("- 所有优化不能破坏已统一的接口。")
    lines.append("")
    lines.append("## 合并原则")
    lines.append("")
    lines.append("- 先统一接口，再分头开发。")
    lines.append("- 先让每个模块能独立运行，再谈整体合并。")
    lines.append("- 合并时优先加适配层，不要直接推翻内部实现。")
    lines.append("- 如果字段或单位要改，必须通知其他三组并同步更新规范。")
    lines.append("")
    lines.append("## AI 开发时必须遵守的事情")
    lines.append("")
    lines.append("- 不得硬编码个人电脑路径。")
    lines.append("- 不得默认某个组的内部实现会永远不变。")
    lines.append("- 不得编造没有确认的硬件参数。")
    lines.append("- 不得把论文参数直接当成项目参数。")
    lines.append("- 不得省略输入输出样例。")
    lines.append("- 不得省略错误样例。")
    lines.append("- 不得不写 README 或运行方式。")
    lines.append("")
    lines.append("## 给 AI 的统一总提示词")
    lines.append("")
    lines.append("```text")
    lines.append("你正在协助焦度计项目的一个独立模块开发。")
    lines.append("请先确认该模块属于 M1 / M2 / M3 / M4 哪一块，再确认当前属于第一阶段、第二阶段还是第三阶段。")
    lines.append("请严格遵守统一接口、统一单位、统一坐标系、统一错误格式、统一日志格式。")
    lines.append("如果成员没有说明板块和阶段，先继续追问，不要擅自假设。")
    lines.append("如果需要其他模块数据，请先使用 mock 数据，不要等待别的模块完成。")
    lines.append("不要硬编码路径，不要改字段名，不要改输出格式。")
    lines.append("请输出：代码、输入样例、输出样例、错误样例、运行方式、当前限制。")
    lines.append("```")
    lines.append("")
    lines.append("## 各板块专用提示词")
    lines.append("")
    lines.append("### M1 输入与配置")
    lines.append("")
    lines.append("```text")
    lines.append("你负责 M1 输入与配置模块。")
    lines.append("只做图片路径、配置文件、任务编号、样本管理和标准输入包输出。")
    lines.append("不要做图像识别，不要做度数计算。")
    lines.append("请输出标准 input_package JSON，并给出错误样例。")
    lines.append("```")
    lines.append("")
    lines.append("### M2 图像识别")
    lines.append("")
    lines.append("```text")
    lines.append("你负责 M2 图像识别模块。")
    lines.append("只做 ROI、去噪、增强、二值化、连通域、质心和 spots JSON。")
    lines.append("不要计算 S/C/A。")
    lines.append("输出要包含识别质量、中间图路径和错误码。")
    lines.append("```")
    lines.append("")
    lines.append("### M3 标定与计算")
    lines.append("")
    lines.append("```text")
    lines.append("你负责 M3 标定与计算模块。")
    lines.append("只接收光斑坐标和配置参数，建立坐标系并输出 S/C/A。")
    lines.append("必须写清楚单位、坐标系和质量判断。")
    lines.append("不确定的公式或硬件参数请标注 TODO_CONFIRM，不要编造。")
    lines.append("```")
    lines.append("")
    lines.append("### M4 本地系统与展示")
    lines.append("")
    lines.append("```text")
    lines.append("你负责 M4 本地系统与展示模块。")
    lines.append("只负责主流程、展示、日志、导出和异常提示。")
    lines.append("第一阶段先用 mock 数据做出可演示的本地系统，不等待其他模块完成。")
    lines.append("不要重写图像识别和度数计算。")
    lines.append("```")
    lines.append("")
    lines.append("## 提交前自查")
    lines.append("")
    lines.append("- 是否独立运行？")
    lines.append("- 是否使用统一 JSON 接口？")
    lines.append("- 是否写清单位和坐标系？")
    lines.append("- 是否输出错误样例？")
    lines.append("- 是否能用 mock 数据跑通？")
    lines.append("- 是否保存日志？")
    lines.append("- 是否提供 README 和运行方式？")
    lines.append("")
    lines.append("## 最终口径")
    lines.append("")
    lines.append("四个模块必须独立开发、统一接口、先做本地系统，再合并成整体，最后再优化成软件。")
    return lines


def build_ai_docx():
    doc = Document()
    setup_doc(doc)
    add_title(doc, "焦度计项目 AI 开发总规范（AI 看版）", "给 AI 开工前读：统一接口、统一阶段、统一合并方式")

    callout(
        doc,
        "开工前必须先问",
        "1. 你负责哪个板块？M1 输入与配置 / M2 图像识别 / M3 标定与计算 / M4 本地系统与展示。\n"
        "2. 你现在处于哪个阶段？第一阶段模块成品 / 第二阶段合并修 bug / 第三阶段优化创新。\n"
        "3. 你希望我先输出什么？代码 / 接口 / 测试样例 / 文档 / 调试建议。\n"
        "4. 你手里现在有哪些本地样本、配置和 mock 数据？\n"
        "如果成员没有回答清楚，AI 不能擅自假设，必须继续追问。"
    )

    heading(doc, "项目目标", 1)
    add_para(doc, "项目目标是把自动焦度计的图片输入变成可读的镜片参数输出。")
    add_para(doc, "输入：标定图 + 测量图 + 配置参数。处理：识别光斑、建立坐标系、计算位移、输出 S/C/A。输出：镜片结果、质量状态、日志、中间图、错误码。")

    heading(doc, "四个独立模块", 1)
    table(doc,
          ["模块", "职责", "输入", "输出", "禁止做什么"],
          [
              ["M1 输入与配置", "管理标定图、测量图、配置文件、样本路径", "图片路径、配置文件、任务编号", "标准输入包 JSON", "不要做图像识别和度数计算"],
              ["M2 图像识别", "预处理、找光斑、算质心", "图片和图像处理参数", "spots JSON", "不要计算 S/C/A"],
              ["M3 标定与计算", "建坐标系、算位移、算 S/C/A", "光斑坐标和光学参数", "result JSON", "不要读取原始图片做识别"],
              ["M4 本地系统与展示", "主流程、展示、日志、导出、异常提示", "各模块输出或 mock 数据", "本地系统成品", "不要重写核心算法"],
          ],
          widths=[1.5, 2.0, 1.65, 1.7, 1.5])

    heading(doc, "必须统一的点", 1)
    table(doc,
          ["统一项", "统一要求"],
          [
              ["输入输出格式", "所有模块都必须使用统一 JSON 结构，字段名不能私改。"],
              ["单位", "像素、毫米、米、角度必须写清楚，不能混用。"],
              ["坐标系", "必须说明原点、X/Y 正方向、图像坐标和测量坐标的关系。"],
              ["错误格式", "错误必须返回统一 error JSON，不能只打印控制台。"],
              ["日志格式", "要保存 task_id、模块名、参数、结果、错误。"],
              ["样本集", "所有组使用同一批样本和 mock 数据进行开发和测试。"],
              ["版本号", "每个模块都要有版本号和变更记录。"],
              ["目录结构", "输入、输出、日志、配置目录要统一命名。"],
              ["运行入口", "每个模块都要能本地独立启动，不能靠手改源码。"],
          ],
          widths=[1.4, 5.4])

    heading(doc, "三个开发阶段", 1)
    heading(doc, "第一阶段：四个模块同时开工", 2)
    add_bullets(doc, [
        "每个模块都必须先做出自己的可运行成品。",
        "允许使用 mock 数据。",
        "允许先做最小可运行版本。",
        "不能等待别的模块完成。",
    ])
    heading(doc, "第二阶段：合并并修复 bug", 2)
    add_bullets(doc, [
        "通过统一接口把四个模块接成完整本地系统。",
        "优先修接口、字段、单位、坐标和错误处理。",
        "不要在这个阶段大范围重构。",
    ])
    heading(doc, "第三阶段：优化创新", 2)
    add_bullets(doc, [
        "在本地系统稳定后做界面、封装、自动导出、自动测试。",
        "后续把本地系统升级成正式软件。",
        "所有优化不能破坏已统一的接口。",
    ])

    heading(doc, "合并原则", 1)
    add_bullets(doc, [
        "先统一接口，再分头开发。",
        "先让每个模块能独立运行，再谈整体合并。",
        "合并时优先加适配层，不要直接推翻内部实现。",
        "如果字段或单位要改，必须通知其他三组并同步更新规范。",
    ])

    heading(doc, "AI 开发时必须遵守的事情", 1)
    add_bullets(doc, [
        "不得硬编码个人电脑路径。",
        "不得默认某个组的内部实现会永远不变。",
        "不得编造没有确认的硬件参数。",
        "不得把论文参数直接当成项目参数。",
        "不得省略输入输出样例。",
        "不得省略错误样例。",
        "不得不写 README 或运行方式。",
    ])

    heading(doc, "给 AI 的统一总提示词", 1)
    code_lines = [
        "你正在协助焦度计项目的一个独立模块开发。",
        "请先确认该模块属于 M1 / M2 / M3 / M4 哪一块，再确认当前属于第一阶段、第二阶段还是第三阶段。",
        "请严格遵守统一接口、统一单位、统一坐标系、统一错误格式、统一日志格式。",
        "如果成员没有说明板块和阶段，先继续追问，不要擅自假设。",
        "如果需要其他模块数据，请先使用 mock 数据，不要等待别的模块完成。",
        "不要硬编码路径，不要改字段名，不要改输出格式。",
        "请输出：代码、输入样例、输出样例、错误样例、运行方式、当前限制。",
    ]
    for line in code_lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=9.5)

    heading(doc, "各板块专用提示词", 1)
    module_prompts = {
        "M1 输入与配置": [
            "你负责 M1 输入与配置模块。",
            "只做图片路径、配置文件、任务编号、样本管理和标准输入包输出。",
            "不要做图像识别，不要做度数计算。",
            "请输出标准 input_package JSON，并给出错误样例。",
        ],
        "M2 图像识别": [
            "你负责 M2 图像识别模块。",
            "只做 ROI、去噪、增强、二值化、连通域、质心和 spots JSON。",
            "不要计算 S/C/A。",
            "输出要包含识别质量、中间图路径和错误码。",
        ],
        "M3 标定与计算": [
            "你负责 M3 标定与计算模块。",
            "只接收光斑坐标和配置参数，建立坐标系并输出 S/C/A。",
            "必须写清楚单位、坐标系和质量判断。",
            "不确定的公式或硬件参数请标注 TODO_CONFIRM，不要编造。",
        ],
        "M4 本地系统与展示": [
            "你负责 M4 本地系统与展示模块。",
            "只负责主流程、展示、日志、导出和异常提示。",
            "第一阶段先用 mock 数据做出可演示的本地系统，不等待其他模块完成。",
            "不要重写图像识别和度数计算。",
        ],
    }
    for title, items in module_prompts.items():
        heading(doc, title, 2)
        for line in items:
            p = doc.add_paragraph()
            r = p.add_run(line)
            set_run_font(r, name="Consolas", size=9.5)

    heading(doc, "提交前自查", 1)
    add_bullets(doc, [
        "是否独立运行？",
        "是否使用统一 JSON 接口？",
        "是否写清单位和坐标系？",
        "是否输出错误样例？",
        "是否能用 mock 数据跑通？",
        "是否保存日志？",
        "是否提供 README 和运行方式？",
    ])

    heading(doc, "最终口径", 1)
    add_para(doc, "四个模块必须独立开发、统一接口、先做本地系统，再合并成整体，最后再优化成软件。")
    doc.save(AI_DOCX)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    human_md = build_human_md()
    ai_md = build_ai_md()
    write_md(HUMAN_MD, human_md)
    write_md(AI_MD, ai_md)
    build_human_docx()
    build_ai_docx()
    print(HUMAN_MD)
    print(HUMAN_DOCX)
    print(AI_MD)
    print(AI_DOCX)


if __name__ == "__main__":
    main()
