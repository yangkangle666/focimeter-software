from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\yangkangle\OneDrive\Desktop\焦度计")
OUT_DIR = ROOT / "outputs"

HUMAN_MD = OUT_DIR / "焦度计项目团队总原则_人看版.md"
HUMAN_DOCX = OUT_DIR / "焦度计项目团队总原则_人看版.docx"
AI_MD = OUT_DIR / "焦度计项目AI开发总规范_AI看版.md"
AI_DOCX = OUT_DIR / "焦度计项目AI开发总规范_AI看版.docx"


def md_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    lines += ["| " + " | ".join(row) + " |" for row in rows]
    return "\n".join(lines)


def set_font(run, size=10.5, bold=False, color="000000", name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def setup_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    for name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 12.5, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        st = doc.styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)


def title(doc, text, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), size=21, bold=True, color="0B2545")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(subtitle), size=11.5, color="666666")
    doc.add_paragraph()


def para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True, color="1F4D78")
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(item))


def h(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def callout(doc, label, text):
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    cell._tc.get_or_add_tcPr().append(shd)
    p = cell.paragraphs[0]
    set_font(p.add_run(label + "："), bold=True, color="1F4D78")
    set_font(p.add_run(text))
    doc.add_paragraph()


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    cells = t.rows[0].cells
    for i, head in enumerate(headers):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E8EEF5")
        cells[i]._tc.get_or_add_tcPr().append(shd)
        p = cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(head), bold=True, color="1F4D78", size=9.5)
        if widths:
            cells[i].width = Inches(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            set_font(p.add_run(str(val)), size=9.2)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return t


def codeblock(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_font(r, size=9.2, name="Consolas", color="333333")


def write(path, lines):
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def human_md_lines():
    module_rows = [
        ["1. 输入与配置模块", "管理标定图、测量图、配置参数、样本路径", "能独立选择图片、读取配置、输出标准数据包", "给图像识别模块提供图片和参数"],
        ["2. 图像识别模块", "ROI、滤波、增强、二值化、连通域、质心提取", "输入图片后输出光斑坐标", "接收输入模块的数据，输出给计算模块"],
        ["3. 标定与计算模块", "建立坐标系、计算光斑位移、输出 S/C/A、质量判断", "输入光斑坐标后输出镜片结果", "接收图像识别结果，输出给系统展示模块"],
        ["4. 本地系统与展示模块", "主界面/主流程、日志、结果展示、报告导出、异常提示", "一个可操作的本地系统外壳", "最后把前三个模块接进来，形成完整系统"],
    ]
    lines = [
        "# 焦度计项目团队总原则（人看版）",
        "",
        "> 给团队成员看的执行规则。重点是四个独立板块、三个开发阶段，以及哪些东西必须统一、绝对不能擅自改。",
        "",
        "## 一句话原则",
        "",
        "四个板块独立开发，互不拖进度；所有板块必须遵守统一接口；第一阶段各自做成品，第二阶段合并修 bug，第三阶段再把本地系统优化成软件。",
        "",
        "## 四大独立板块",
        "",
        md_table(["板块", "主要功能", "独立成品", "和其他板块的联系"], module_rows),
        "",
        "### 四板块执行要求",
        "",
        "- 每个板块必须能单独运行，不能等其他板块完成才开始。",
        "- 每个板块必须有自己的输入样例、输出样例、错误样例。",
        "- 每个板块可以先用模拟数据开发，但输出格式必须和正式接口一致。",
        "- 板块内部怎么实现可以不同，但对外输出必须统一。",
        "- 合并时只通过接口合并，不允许临时手工复制粘贴数据。",
        "",
        "## 三个开发阶段",
        "",
        md_table(
            ["阶段", "目标", "四组怎么做", "验收标准"],
            [
                ["第一阶段：四个板块同时开工", "各自做出能运行的模块成品", "每组独立开发；允许使用 mock 数据；不能等待其他组", "每组能独立演示自己的模块"],
                ["第二阶段：四个板块合并修 bug", "把四个模块接成完整本地系统", "按接口依次接入；优先修字段、单位、坐标、错误处理", "本地系统能跑完整流程"],
                ["第三阶段：优化创新", "把本地系统逐步变成正式软件", "做界面、封装、报告、批量测试、相机接入等", "系统更稳定、更易用、更像软件"],
            ],
        ),
        "",
        "## 必须统一到什么程度",
        "",
        md_table(
            ["统一项", "统一到什么程度", "绝不允许"],
            [
                ["接口格式", "所有模块统一用 JSON 输入输出，字段名固定", "不能私自改字段名或层级"],
                ["图片路径", "统一使用相对路径或配置路径", "不能写死个人电脑绝对路径"],
                ["单位", "图像坐标用 pixel；像元尺寸用 um；长度用 m/mm 且字段名写单位；角度用 degree；屈光度用 D", "不能像素、毫米、米混着算"],
                ["坐标系", "图像坐标原点左上角；X 向右，Y 向下；标定坐标以中心光斑为原点", "不能不说明坐标类型就输出 x/y"],
                ["光斑输出", "统一输出 spot_id、role、x、y、confidence、status", "不能只打印控制台结果"],
                ["计算输出", "统一输出 S、C、A、unit、lens_type、quality、error", "不能只输出一个数字"],
                ["错误格式", "统一 error.code、error.message、error.module、recoverable", "不能程序直接崩溃不返回原因"],
                ["日志格式", "统一记录 task_id、module、start_time、end_time、status、input、output、error", "不能没有运行记录"],
                ["版本记录", "每个模块写版本号、改动内容、当前问题", "不能无法追踪谁改了什么"],
            ],
        ),
        "",
        "## 绝对不能擅自改动",
        "",
        "- 不准擅自改四大板块名称和职责。",
        "- 不准擅自改统一字段名。",
        "- 不准擅自改单位。",
        "- 不准擅自改坐标系定义。",
        "- 不准擅自改错误码和日志格式。",
        "- 不准把个人电脑路径写进正式代码。",
        "- 不准因为自己方便就绕过接口。",
        "- 不准在第二阶段合并时临时推翻其他组成果。",
        "",
        "## 合并方式",
        "",
        "合并不是把四组代码硬拼到一起，而是本地系统主流程按接口调用四个模块：",
        "",
        "```text",
        "输入与配置模块 → 图像识别模块 → 标定与计算模块 → 本地系统与展示模块",
        "```",
        "",
        "如果某个模块暂时没完成，其他模块用 mock 数据继续推进，不能停工等待。",
    ]
    return lines


def ai_md_lines():
    lines = [
        "# 焦度计项目 AI 开发总规范（AI 看版）",
        "",
        "> 给 AI 开工前阅读。AI 必须先确认成员负责哪个板块，再按该板块规则开发。本文直接规定统一格式，不只说“要统一”。",
        "",
        "## AI 开工前必须先问成员",
        "",
        "在写任何代码或文档前，AI 必须先问：",
        "",
        "1. 你负责哪个板块？只能选择：`M1 输入与配置模块`、`M2 图像识别模块`、`M3 标定与计算模块`、`M4 本地系统与展示模块`。",
        "2. 当前处于哪个阶段？只能选择：`第一阶段模块成品`、`第二阶段合并修 bug`、`第三阶段优化创新`。",
        "3. 你要我先做什么？代码、接口、测试样例、README、调试、报告，还是其他？",
        "4. 你现在有哪些本地文件？标定图、测量图、配置文件、mock 数据分别在哪里？",
        "",
        "如果成员没有回答清楚，AI 必须继续追问，不能直接开始。",
        "",
        "## 四大模块边界",
        "",
        md_table(
            ["模块", "允许做", "必须输出", "禁止做"],
            [
                ["M1 输入与配置", "读取图片路径、配置参数、样本列表、任务编号", "标准输入数据包 input_package.json", "禁止做图像识别、禁止计算 S/C/A"],
                ["M2 图像识别", "ROI、滤波、增强、二值化、连通域、质心", "spots_calib.json 和 spots_meas.json", "禁止计算镜片度数"],
                ["M3 标定与计算", "坐标系、光斑位移、镜片类型、S/C/A、质量判断", "result.json", "禁止读取原始图片重新识别"],
                ["M4 本地系统与展示", "主流程、界面/命令行、日志、结果展示、导出、异常提示", "可操作本地系统", "禁止重写 M2/M3 核心算法"],
            ],
        ),
        "",
        "## 统一目录结构",
        "",
        "建议所有本地模块按下面结构组织，至少目录名称要统一：",
        "",
        "```text",
        "focimeter_system/",
        "  config/default_config.json",
        "  data/samples/calibration/",
        "  data/samples/measurement/",
        "  data/mock/",
        "  modules/input_config/",
        "  modules/image_recognition/",
        "  modules/calibration_calculation/",
        "  modules/local_system/",
        "  outputs/images/",
        "  outputs/results/",
        "  outputs/logs/",
        "  outputs/reports/",
        "```",
        "",
        "## 统一输入数据包",
        "",
        "所有模块接收任务时，统一使用这个输入包结构：",
        "",
        "```json",
        "{",
        "  \"task_id\": \"sample_001\",",
        "  \"calibration_image\": \"data/samples/calibration/calib_001.png\",",
        "  \"measurement_image\": \"data/samples/measurement/meas_001.png\",",
        "  \"config_path\": \"config/default_config.json\",",
        "  \"run_mode\": \"local_image\"",
        "}",
        "```",
        "",
        "字段不可改名。路径优先使用相对路径，不能硬编码个人电脑绝对路径。",
        "",
        "## 统一配置文件",
        "",
        "配置统一使用 `config/default_config.json`，结构如下：",
        "",
        "```json",
        "{",
        "  \"camera\": {",
        "    \"pixel_size_um\": 4.0,",
        "    \"image_width\": null,",
        "    \"image_height\": null",
        "  },",
        "  \"optical\": {",
        "    \"distance_m\": 0.03,",
        "    \"hartmann_spacing_mm\": null",
        "  },",
        "  \"image_processing\": {",
        "    \"roi_width_ratio\": 0.9,",
        "    \"roi_height_ratio\": 0.9,",
        "    \"median_kernel\": 3,",
        "    \"tophat_kernel\": 30,",
        "    \"otsu_a\": 0.4,",
        "    \"otsu_b\": 0.7,",
        "    \"max_depth\": 2",
        "  },",
        "  \"recognition\": {",
        "    \"expected_spot_count\": 5,",
        "    \"min_confidence\": 0.7",
        "  },",
        "  \"calculation\": {",
        "    \"pixel_threshold\": 1.0,",
        "    \"angle_unit\": \"degree\"",
        "  }",
        "}",
        "```",
        "",
        "不知道的真实硬件参数写 `null` 或 `TODO_CONFIRM`，不能编造。",
        "",
        "## 统一光斑识别输出",
        "",
        "M2 必须输出如下格式：",
        "",
        "```json",
        "{",
        "  \"task_id\": \"sample_001\",",
        "  \"status\": \"ok\",",
        "  \"image_type\": \"calibration\",",
        "  \"coordinate_type\": \"image_pixel\",",
        "  \"spots\": [",
        "    {",
        "      \"spot_id\": 0,",
        "      \"role\": \"center\",",
        "      \"x\": 512.34,",
        "      \"y\": 384.21,",
        "      \"confidence\": 0.96",
        "    }",
        "  ],",
        "  \"quality\": {",
        "    \"expected_count\": 5,",
        "    \"detected_count\": 5,",
        "    \"is_usable\": true",
        "  },",
        "  \"error\": null",
        "}",
        "```",
        "",
        "统一 `role` 值：`center`、`y_positive`、`x_positive`、`left_or_negative`、`other`。如果不能判断角色，必须写 `unknown` 并给 warning。",
        "",
        "## 统一计算输出",
        "",
        "M3 必须输出如下格式：",
        "",
        "```json",
        "{",
        "  \"task_id\": \"sample_001\",",
        "  \"status\": \"ok\",",
        "  \"lens_type\": \"spherical\",",
        "  \"result\": {",
        "    \"S\": -2.50,",
        "    \"C\": 0.00,",
        "    \"A\": null,",
        "    \"unit\": \"D\"",
        "  },",
        "  \"quality\": {",
        "    \"is_usable\": true,",
        "    \"confidence\": 0.91,",
        "    \"warnings\": []",
        "  },",
        "  \"intermediate\": {",
        "    \"coordinate_system_valid\": true,",
        "    \"shift_unit\": \"pixel\"",
        "  },",
        "  \"error\": null",
        "}",
        "```",
        "",
        "## 统一错误格式",
        "",
        "所有模块错误必须返回：",
        "",
        "```json",
        "{",
        "  \"status\": \"error\",",
        "  \"error\": {",
        "    \"code\": \"SPOT_COUNT_MISMATCH\",",
        "    \"message\": \"Expected 5 spots but detected 3.\",",
        "    \"module\": \"image_recognition\",",
        "    \"recoverable\": true",
        "  }",
        "}",
        "```",
        "",
        "统一错误码：`IMAGE_NOT_FOUND`、`CONFIG_NOT_FOUND`、`CONFIG_INVALID`、`IMAGE_LOAD_FAILED`、`SPOT_COUNT_MISMATCH`、`CENTROID_FAILED`、`COORDINATE_SYSTEM_INVALID`、`UNIT_MISMATCH`、`CALCULATION_FAILED`、`UNKNOWN_ERROR`。",
        "",
        "## 统一单位",
        "",
        md_table(
            ["物理量", "统一单位", "字段写法"],
            [
                ["图像坐标", "pixel", "`x`、`y`，并写 `coordinate_type=image_pixel`"],
                ["像元尺寸", "um", "`pixel_size_um`"],
                ["物理距离", "m", "`distance_m`"],
                ["光阑间距", "mm", "`hartmann_spacing_mm`"],
                ["角度", "degree", "`A`、`angle_unit=degree`"],
                ["屈光度", "D", "`S`、`C`、`unit=D`"],
                ["置信度", "0 到 1", "`confidence=0.91`，不用百分数"],
            ],
        ),
        "",
        "## 统一坐标系",
        "",
        "图像坐标：左上角为原点，X 向右为正，Y 向下为正，单位 pixel。",
        "",
        "标定坐标：中心光斑为原点，`y_positive` 定义 Y 轴正方向，`x_positive` 定义 X 轴正方向。M3 计算前必须说明坐标是否已经从图像坐标转换到标定坐标。",
        "",
        "## 统一日志格式",
        "",
        "每次运行至少保存：",
        "",
        "```json",
        "{",
        "  \"task_id\": \"sample_001\",",
        "  \"module\": \"image_recognition\",",
        "  \"start_time\": \"2026-07-18 10:00:00\",",
        "  \"end_time\": \"2026-07-18 10:00:01\",",
        "  \"status\": \"ok\",",
        "  \"input_files\": [],",
        "  \"output_files\": [],",
        "  \"parameters\": {},",
        "  \"warnings\": [],",
        "  \"error\": null",
        "}",
        "```",
        "",
        "## AI 开发时的硬性要求",
        "",
        "- 先问成员负责哪个板块，再开始。",
        "- 先问当前阶段，再开始。",
        "- 需要其他模块时，先用 mock 数据。",
        "- 不能擅自改字段名、单位、坐标系、错误码。",
        "- 不确定的硬件参数必须标注 `TODO_CONFIRM`。",
        "- 每次输出都要包含运行方法、输入样例、输出样例、错误样例。",
    ]
    return lines


def build_human_docx(lines):
    doc = Document()
    setup_doc(doc)
    title(doc, "焦度计项目团队总原则（人看版）", "四个独立板块、三个开发阶段、统一规范和禁止擅改事项")
    callout(doc, "一句话原则", "四个板块独立开发，互不拖进度；所有板块必须遵守统一接口；第一阶段各自做成品，第二阶段合并修 bug，第三阶段再把本地系统优化成软件。")
    h(doc, "四大独立板块", 1)
    table(doc, ["板块", "主要功能", "独立成品", "和其他板块的联系"], [
        ["1. 输入与配置模块", "管理标定图、测量图、配置参数、样本路径", "能独立选择图片、读取配置、输出标准数据包", "给图像识别模块提供图片和参数"],
        ["2. 图像识别模块", "ROI、滤波、增强、二值化、连通域、质心提取", "输入图片后输出光斑坐标", "接收输入模块的数据，输出给计算模块"],
        ["3. 标定与计算模块", "建立坐标系、计算光斑位移、输出 S/C/A、质量判断", "输入光斑坐标后输出镜片结果", "接收图像识别结果，输出给系统展示模块"],
        ["4. 本地系统与展示模块", "主界面/主流程、日志、结果展示、报告导出、异常提示", "一个可操作的本地系统外壳", "最后把前三个模块接进来，形成完整系统"],
    ], widths=[1.35, 2.15, 2.1, 2.0])
    h(doc, "四板块执行要求", 1)
    bullets(doc, [
        "每个板块必须能单独运行，不能等其他板块完成才开始。",
        "每个板块必须有自己的输入样例、输出样例、错误样例。",
        "每个板块可以先用模拟数据开发，但输出格式必须和正式接口一致。",
        "板块内部怎么实现可以不同，但对外输出必须统一。",
        "合并时只通过接口合并，不允许临时手工复制粘贴数据。",
    ])
    h(doc, "三个开发阶段", 1)
    table(doc, ["阶段", "目标", "四组怎么做", "验收标准"], [
        ["第一阶段：四个板块同时开工", "各自做出能运行的模块成品", "每组独立开发；允许使用 mock 数据；不能等待其他组", "每组能独立演示自己的模块"],
        ["第二阶段：四个板块合并修 bug", "把四个模块接成完整本地系统", "按接口依次接入；优先修字段、单位、坐标、错误处理", "本地系统能跑完整流程"],
        ["第三阶段：优化创新", "把本地系统逐步变成正式软件", "做界面、封装、报告、批量测试、相机接入等", "系统更稳定、更易用、更像软件"],
    ], widths=[1.8, 2.0, 2.45, 1.55])
    h(doc, "必须统一到什么程度", 1)
    table(doc, ["统一项", "统一到什么程度", "绝不允许"], [
        ["接口格式", "所有模块统一用 JSON 输入输出，字段名固定", "不能私自改字段名或层级"],
        ["图片路径", "统一使用相对路径或配置路径", "不能写死个人电脑绝对路径"],
        ["单位", "图像坐标 pixel；像元 um；长度 m/mm；角度 degree；屈光度 D", "不能像素、毫米、米混着算"],
        ["坐标系", "图像坐标原点左上角；标定坐标以中心光斑为原点", "不能不说明坐标类型就输出 x/y"],
        ["光斑输出", "统一输出 spot_id、role、x、y、confidence、status", "不能只打印控制台结果"],
        ["计算输出", "统一输出 S、C、A、unit、lens_type、quality、error", "不能只输出一个数字"],
        ["错误格式", "统一 error.code、error.message、error.module、recoverable", "不能程序直接崩溃不返回原因"],
        ["日志格式", "统一记录 task_id、module、start_time、end_time、status、input、output、error", "不能没有运行记录"],
    ], widths=[1.25, 4.0, 2.1])
    h(doc, "绝对不能擅自改动", 1)
    bullets(doc, [
        "不准擅自改四大板块名称和职责。",
        "不准擅自改统一字段名、单位、坐标系定义、错误码和日志格式。",
        "不准把个人电脑路径写进正式代码。",
        "不准因为自己方便就绕过接口。",
        "不准在第二阶段合并时临时推翻其他组成果。",
    ])
    h(doc, "合并方式", 1)
    para(doc, "合并不是把四组代码硬拼到一起，而是本地系统主流程按接口调用四个模块：输入与配置模块 → 图像识别模块 → 标定与计算模块 → 本地系统与展示模块。")
    para(doc, "如果某个模块暂时没完成，其他模块用 mock 数据继续推进，不能停工等待。")
    doc.save(HUMAN_DOCX)


def build_ai_docx(lines):
    doc = Document()
    setup_doc(doc)
    title(doc, "焦度计项目 AI 开发总规范（AI 看版）", "先确认板块和阶段，再按统一接口开发")
    callout(doc, "AI 开工前必须先问", "1. 你负责哪个板块？M1 输入与配置 / M2 图像识别 / M3 标定与计算 / M4 本地系统与展示。\n2. 当前处于哪个阶段？第一阶段模块成品 / 第二阶段合并修 bug / 第三阶段优化创新。\n3. 你要我先做什么？代码、接口、测试样例、README、调试、报告，还是其他？\n4. 你现在有哪些本地文件？标定图、测量图、配置文件、mock 数据分别在哪里？")
    h(doc, "四大模块边界", 1)
    table(doc, ["模块", "允许做", "必须输出", "禁止做"], [
        ["M1 输入与配置", "读取图片路径、配置参数、样本列表、任务编号", "标准输入数据包 input_package.json", "禁止做图像识别、禁止计算 S/C/A"],
        ["M2 图像识别", "ROI、滤波、增强、二值化、连通域、质心", "spots_calib.json 和 spots_meas.json", "禁止计算镜片度数"],
        ["M3 标定与计算", "坐标系、光斑位移、镜片类型、S/C/A、质量判断", "result.json", "禁止读取原始图片重新识别"],
        ["M4 本地系统与展示", "主流程、界面/命令行、日志、结果展示、导出、异常提示", "可操作本地系统", "禁止重写 M2/M3 核心算法"],
    ], widths=[1.5, 2.45, 1.9, 2.1])
    h(doc, "统一目录结构", 1)
    codeblock(doc, [
        "focimeter_system/",
        "  config/default_config.json",
        "  data/samples/calibration/",
        "  data/samples/measurement/",
        "  data/mock/",
        "  modules/input_config/",
        "  modules/image_recognition/",
        "  modules/calibration_calculation/",
        "  modules/local_system/",
        "  outputs/images/ outputs/results/ outputs/logs/ outputs/reports/",
    ])
    h(doc, "统一输入数据包", 1)
    codeblock(doc, [
        "{",
        "  \"task_id\": \"sample_001\",",
        "  \"calibration_image\": \"data/samples/calibration/calib_001.png\",",
        "  \"measurement_image\": \"data/samples/measurement/meas_001.png\",",
        "  \"config_path\": \"config/default_config.json\",",
        "  \"run_mode\": \"local_image\"",
        "}",
    ])
    h(doc, "统一配置文件", 1)
    codeblock(doc, [
        "{",
        "  \"camera\": {\"pixel_size_um\": 4.0, \"image_width\": null, \"image_height\": null},",
        "  \"optical\": {\"distance_m\": 0.03, \"hartmann_spacing_mm\": null},",
        "  \"image_processing\": {\"roi_width_ratio\": 0.9, \"roi_height_ratio\": 0.9, \"median_kernel\": 3, \"tophat_kernel\": 30, \"otsu_a\": 0.4, \"otsu_b\": 0.7, \"max_depth\": 2},",
        "  \"recognition\": {\"expected_spot_count\": 5, \"min_confidence\": 0.7},",
        "  \"calculation\": {\"pixel_threshold\": 1.0, \"angle_unit\": \"degree\"}",
        "}",
    ])
    h(doc, "统一光斑识别输出", 1)
    codeblock(doc, [
        "{",
        "  \"task_id\": \"sample_001\", \"status\": \"ok\",",
        "  \"image_type\": \"calibration\", \"coordinate_type\": \"image_pixel\",",
        "  \"spots\": [{\"spot_id\": 0, \"role\": \"center\", \"x\": 512.34, \"y\": 384.21, \"confidence\": 0.96}],",
        "  \"quality\": {\"expected_count\": 5, \"detected_count\": 5, \"is_usable\": true},",
        "  \"error\": null",
        "}",
    ])
    para(doc, "统一 role 值：center、y_positive、x_positive、left_or_negative、other、unknown。")
    h(doc, "统一计算输出", 1)
    codeblock(doc, [
        "{",
        "  \"task_id\": \"sample_001\", \"status\": \"ok\", \"lens_type\": \"spherical\",",
        "  \"result\": {\"S\": -2.50, \"C\": 0.00, \"A\": null, \"unit\": \"D\"},",
        "  \"quality\": {\"is_usable\": true, \"confidence\": 0.91, \"warnings\": []},",
        "  \"intermediate\": {\"coordinate_system_valid\": true, \"shift_unit\": \"pixel\"},",
        "  \"error\": null",
        "}",
    ])
    h(doc, "统一错误格式", 1)
    codeblock(doc, [
        "{",
        "  \"status\": \"error\",",
        "  \"error\": {\"code\": \"SPOT_COUNT_MISMATCH\", \"message\": \"Expected 5 spots but detected 3.\", \"module\": \"image_recognition\", \"recoverable\": true}",
        "}",
    ])
    h(doc, "统一单位", 1)
    table(doc, ["物理量", "统一单位", "字段写法"], [
        ["图像坐标", "pixel", "x、y，且 coordinate_type=image_pixel"],
        ["像元尺寸", "um", "pixel_size_um"],
        ["物理距离", "m", "distance_m"],
        ["光阑间距", "mm", "hartmann_spacing_mm"],
        ["角度", "degree", "A、angle_unit=degree"],
        ["屈光度", "D", "S、C、unit=D"],
        ["置信度", "0 到 1", "confidence=0.91，不用百分数"],
    ], widths=[1.4, 1.2, 4.5])
    h(doc, "统一坐标系", 1)
    para(doc, "图像坐标：左上角为原点，X 向右为正，Y 向下为正，单位 pixel。")
    para(doc, "标定坐标：中心光斑为原点，y_positive 定义 Y 轴正方向，x_positive 定义 X 轴正方向。M3 计算前必须说明坐标是否已经从图像坐标转换到标定坐标。")
    h(doc, "AI 开发硬性要求", 1)
    bullets(doc, [
        "先问成员负责哪个板块，再开始。",
        "先问当前阶段，再开始。",
        "需要其他模块时，先用 mock 数据。",
        "不能擅自改字段名、单位、坐标系、错误码。",
        "不确定的硬件参数必须标注 TODO_CONFIRM。",
        "每次输出都要包含运行方法、输入样例、输出样例、错误样例。",
    ])
    doc.save(AI_DOCX)


def main():
    OUT_DIR.mkdir(exist_ok=True)
    human = human_md_lines()
    ai = ai_md_lines()
    write(HUMAN_MD, human)
    write(AI_MD, ai)
    build_human_docx(human)
    build_ai_docx(ai)
    print(HUMAN_MD)
    print(HUMAN_DOCX)
    print(AI_MD)
    print(AI_DOCX)


if __name__ == "__main__":
    main()
